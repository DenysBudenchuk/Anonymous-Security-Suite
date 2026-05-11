#!/usr/bin/env python3
"""
Polish Document Anonymizer (Offline Secure Edition)
==========================
Wymagania:
    pip install spacy transformers torch pymupdf huggingface_hub python-docx watchdog pystray pillow opencv-python-headless numpy
"""

import os
import sys
import types
from typing import Optional
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import threading
import re
import json
from pathlib import Path
from datetime import datetime
import time
import queue
import gc

import cv2
import numpy as np
from PIL import ImageTk, Image, ImageDraw
import pystray
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

# ─────────────────────────────────────────────
# Змінні глобальні
# ─────────────────────────────────────────────
spacy = None
nlp_spacy = None
pipeline_pii = None
fitz: types.ModuleType | None = None   # FIX: анотація типу замість голого None

MODELS_LOADED = False
# ВИДАЛЕНО: MODELS_LOADING — ніде не читалась

if getattr(sys, 'frozen', False):
    application_path = Path(sys.executable).parent
    EXE_PATH = os.path.abspath(sys.executable)
else:
    application_path = Path(__file__).parent
    EXE_PATH = os.path.abspath(__file__)

BASE_MODEL_PATH = application_path / "offline_models"

# ==========================================
# СТАТИЧНА БАЗА ДАНИХ (GAZETTEER)
# ==========================================
KNOWN_NAMES_SET: set[str] = set()

STOP_WORDS = {
    # 1. Займенники, сполучники, прийменники
    "ja", "ty", "on", "ona", "ono", "my", "wy", "oni", "one", "mnie", "mi", "mną", "ciebie", "cię", "ci",
    "tobą", "jego", "niego", "go", "jemu", "niemu", "mu", "nim", "jej", "niej", "nią", "nas", "nam", "nami",
    "was", "wam", "wami", "ich", "nich", "im", "nimi", "siebie", "się", "sobie", "sobą", "mój", "moja", "moje",
    "twój", "twoja", "twoje", "nasz", "nasza", "nasze", "wasz", "wasza", "wasze", "swój", "swoja", "swoje",
    "ten", "ta", "to", "tamten", "tamta", "tamto", "taki", "taka", "takie", "tym", "obie", "oba", "oboje",
    "kto", "co", "jaki", "który", "ktoś", "coś", "jakiś", "nikt", "nic", "żaden", "każdy", "wszyscy", "wielu",
    "w", "we", "z", "ze", "na", "do", "dla", "o", "od", "po", "pod", "pode", "nad", "nade", "przed", "przede",
    "za", "przy", "bez", "beze", "u", "ku", "mimo", "pomimo", "wokół", "obok", "wzdłuż", "oprócz", "zamiast",
    "według", "wobec", "pomiędzy", "między", "wewnątrz", "zewnątrz", "poniżej", "powyżej",
    "i", "a", "oraz", "lub", "czy", "albo", "bądź", "ani", "ni", "ale", "lecz", "jednak", "jednakże",
    "zaś", "natomiast", "więc", "dlatego", "zatem", "toteż", "że", "iż", "ponieważ", "bo", "gdy", "gdyby",
    "jeśli", "jeżeli", "choć", "chociaż", "chociażby", "zanim", "aż", "nie", "tak", "niech", "by", "no",
    "oby", "tylko", "nawet", "też", "także", "ponadto", "stąd", "wreszcie", "następnie", "obecnie", "aktualnie",

    # 2. Дієслова та прислівники
    "być", "jest", "są", "był", "była", "było", "byli", "były", "będzie", "będą", "będę", "będziesz",
    "mieć", "ma", "mają", "miał", "miała", "miało", "mieli", "miały", "mam", "masz",
    "może", "mogą", "mógł", "mogła", "musieć", "musi", "muszą", "chcieć", "chce", "chcą",
    "został", "została", "zostali", "zostały", "zostać", "powinien", "powinna", "powinno", "powinni",
    "nierzetelnie", "niezwłocznie", "szybko", "wolno", "bardzo", "mało", "dużo", "wiele", "nieco", "zbyt",

    # 3. Корпоративна лексика, посади, титули
    "pan", "pana", "panu", "panem", "pani", "panią", "państwo", "państwa", "państwu", "proszę", "dziękuję",
    "witam", "pozdrawiam", "szanowny", "szanowna", "szanowni", "poważaniem", "poważanie",
    "dnia", "roku", "miesiąca", "ulica", "numer", "telefon", "faks", "email", "mail", "adres",
    "miasto", "miejscowość", "kod", "pocztowy", "nip", "pesel", "regon", "iban", "krs", "cvv", "cvc", "bdo",
    "dowód", "osobisty", "paszport", "załącznik", "dokument", "umowa", "faktura", "paragon", "rachunek",
    "kwota", "suma", "cena", "netto", "brutto", "podatek", "vat", "pit", "cit", "data", "podpis", "pieczęć",
    "sztuka", "uwaga", "strona", "konto", "karta", "kredytowa", "debetowa", "bank", "przelew", "gotówka",
    "imię", "nazwisko", "dane", "osoba", "excel", "dz", "chaos", "wartość", "koszt", "przychód", "dochód",
    "marszałek", "marszałka", "marszałkowi", "marszałkiem", "marszałkowskiego",
    "konserwator", "konserwatora", "konserwatorem",
    "inżynier", "inżyniera", "inżynierem", "dyrektor", "dyrektora", "dyrektorze", "dyrekcja", "główny", "głównego",
    "kierownik", "kierownika", "prezes", "prezesa", "wiceprezes", "członek", "zastępca", "pełnomocnik",
    "wójt", "burmistrz", "starosta", "wojewoda", "poseł", "senator", "sędzia", "prokurator", "adwokat", "radca",
    "delegatura", "delegatury", "zarządzenie", "zarządzenia", "uchwała", "uchwały", "ustawa", "ustawy",
    "wieloletniej", "prognozie", "finansowej", "rozporządzenie", "decyzja", "postanowienie", "zaświadczenie",

    # 4. Установи, організації та юридичні терміни
    "firma", "spółka", "spółki", "spółką", "przedsiębiorstwo", "działalność", "gospodarcza",
    "urząd", "urzędu", "urzędzie", "urzędem", "gmina", "powiat", "powiatowy",
    "wydział", "wydziału", "departament", "departamentu", "departamencie", "referat", "sekcja",
    "ministerstwo", "ministerstwa", "minister", "ministra", "ministrem", "ministrowi",
    "komisja", "komisji", "komisję", "związek", "związku", "związkiem", "rada", "zarząd",
    "teatr", "teatru", "teatrowi", "teatrem", "teatrze", "muzeum", "szkoła", "uczelnia", "szpital", "przychodnia",
    "województwo", "województwa", "województw", "wojewódzkiego", "wojewódzkiej", "sejmik", "sejmiku",
    "sąd", "sądu", "sądzie", "trybunał", "prokuratura", "policja", "straż",

    # 5. Міста, країни, географія
    "kielce", "kielcach", "kielcami", "kielc", "kieleckich", "kieleckiej", "kiele", "kie",
    "warszawa", "warszawie", "łódź", "łodzi", "częstochowa", "częstochowy", "kraków", "krakowie",
    "poznań", "poznaniu", "wrocław", "wrocławiu", "gdańsk", "gdańsku", "szczecin", "szczecinie",
    "bydgoszcz", "bydgoszczy", "lublin", "lublinie", "katowice", "katowicach", "białystok", "białymstoku",
    "gdynia", "sopot", "rzeszów", "rzeszowie", "toruń", "toruniu", "opole", "opolu",
    "europa", "europie", "islandia", "islandii", "islandzką", "norwegia", "norwegii", "ukraina", "ukrainę", "ue",

    # 6. Абревіатури, одиниці виміру та OCR-сміття
    # ВИДАЛЕНО: однобуквені "a".."z" — ніколи не спрацьовували (regex вимагає Uppercase першої букви)
    "ul", "al", "pl", "nr", "tel", "fax", "kom", "km", "cm", "mm", "kg", "szt", "egz", "kpl",
    "godz", "min", "sek", "sp", "zoo", "sa", "prof", "dr", "mgr", "inż", "lek", "med", "hab",
    "np", "tj", "tzw", "ww", "itd", "itp", "cdn", "str", "poz", "ust", "art", "par", "pkt",
    "im", "ce", "cach", "ręb", "krzy", "święt", "ill", "nro", "uzp", "inw", "ke",

    # 7. Римські цифри (часто сприймаються як імена)
    "ii", "iii", "iv", "vi", "vii", "viii", "ix", "xi", "xii", "xiii", "xiv", "xv",

    # 8. Державні прикметники
    "polska", "polskiej", "polskich", "polską", "polski", "polskie",
    "rzeczypospolita", "rzeczypospolitej", "unijny", "unijna", "unijne", "unijnych",
    "europejski", "europejska", "europejskie", "europejskiego", "krajowy", "krajowa", "krajowe",

    # 9. Звичайні слова, що спричинили False Positives
    "poprawa", "stan", "zasady", "prawna", "łącza", "łączna", "konieczne", "ośmiomiesięcznego",

    # 10. Інші загальні іменники / прикметники
    "mucha", "zima", "lato", "wiosna", "jesień", "kowal", "kruk", "lis", "wilk", "niedziela", "sobota",
    "zając", "niedźwiedź", "sowa", "dudek", "kaczka", "gęś", "ptak", "ryba", "kot", "pies",
    "cebula", "burak", "kapusta", "marchew", "woda", "piwo", "wino", "chleb", "masło", "ser",
    "wiatr", "mróz", "burza", "chmura", "słońce", "księżyc", "gwiazda", "niebo", "ziemia",
    "biały", "czarny", "czerwony", "zielony", "niebieski", "żółty", "szary", "brązowy", "złoty", "srebrny",
    "mały", "duży", "gruby", "chudy", "wysoki", "niski", "stary", "nowy", "młody",
    "dobry", "zły", "prawy", "lewy", "góra", "dół", "przód", "tył", "bok", "środek",
    "koniec", "początek", "wielki", "krótki", "długi", "las", "pole", "morze", "rzeka", "góry",
    "dzisiaj", "jutro", "wczoraj", "rano", "wieczór", "południe", "kwartał", "półrocze", "termin", "okres",
}


def load_names_db(log_callback=None):
    """Завантажує базу імен у геш-множину для швидкого пошуку O(1)."""
    global KNOWN_NAMES_SET
    db_path = BASE_MODEL_PATH / "names_db.txt"

    if db_path.exists():
        try:
            with open(db_path, "r", encoding="utf-8") as f:
                KNOWN_NAMES_SET.update(line.strip().lower() for line in f if line.strip())
            if log_callback:
                log_callback(f"✅ Baza słownikowa załadowana: {len(KNOWN_NAMES_SET)} słów.")
        except Exception as e:
            if log_callback:
                log_callback(f"⚠️ Błąd ładowania bazy imion: {e}")
    else:
        if log_callback:
            log_callback("⚠️ Brak pliku names_db.txt w folderze offline_models.")


# ─────────────────────────────────────────────
# Ініціалізація моделей і OCR
# ─────────────────────────────────────────────
def ensure_models_exist(log_callback=None):
    # ВИДАЛЕНО: перенаправлення sys.stderr/sys.stdout — log_callback вже перехоплює все

    os.environ['CURL_CA_BUNDLE'] = ''
    os.environ['PYTHONHTTPSVERIFY'] = '0'
    os.environ["HF_HUB_DISABLE_PROGRESS_BARS"] = "1"

    path_pii = BASE_MODEL_PATH / "eu-pii-anonimization"
    path_spacy = BASE_MODEL_PATH / "pl_core_news_lg"
    path_tesseract = BASE_MODEL_PATH / "tesseract"
    tesseract_exe = path_tesseract / "tesseract.exe"

    pii_missing = not path_pii.exists() or not list(path_pii.glob('*'))
    spacy_missing = not path_spacy.exists() or not list(path_spacy.glob('*'))

    if pii_missing or spacy_missing:
        if log_callback:
            log_callback("⚠️ Brak modeli AI. Rozpoczynam pobieranie...")
        BASE_MODEL_PATH.mkdir(parents=True, exist_ok=True)

        from huggingface_hub import snapshot_download
        from huggingface_hub.utils.tqdm import disable_progress_bars  # FIX: правильний імпорт
        disable_progress_bars()

        HF_TOKEN = "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"

        try:
            if pii_missing:
                if log_callback:
                    log_callback("⏳ Pobieranie PII (Transformers)...")
                snapshot_download(  # type: ignore[call-overload]
                    repo_id="bardsai/eu-pii-anonimization-multilang",
                    local_dir=str(path_pii),
                    token=HF_TOKEN,
                    local_dir_use_symlinks=False,
                    ignore_patterns=["*.msgpack", "*.h5", "*.ot", "*.onnx", "*.flax"],
                )

            if spacy_missing:
                if log_callback:
                    log_callback("⏳ Pobieranie spaCy...")
                snapshot_download(  # type: ignore[call-overload]
                    repo_id="spacy/pl_core_news_lg",
                    local_dir=str(path_spacy),
                    token=HF_TOKEN,
                    local_dir_use_symlinks=False,
                    ignore_patterns=["*.h5", "*.ot", "*.onnx", "*.flax"],
                )
            if log_callback:
                log_callback("✅ Modele AI pobrane.")
        except Exception as e:
            if log_callback:
                log_callback(f"❌ Błąd pobierania modeli AI: {e}")
            raise

    if tesseract_exe.exists():
        abs_tess_path = str(path_tesseract.absolute())
        abs_tessdata_path = str((path_tesseract / "tessdata").absolute())
        os.environ["PATH"] = abs_tess_path + os.pathsep + os.environ.get("PATH", "")
        os.environ["TESSDATA_PREFIX"] = abs_tessdata_path
        if log_callback:
            log_callback("✅ Silnik OCR aktywny (Ścieżki bezwzględne).")
    else:
        if log_callback:
            log_callback("⚠️ Brak silnika OCR - skany nie będą obsługiwane.")


def lazy_load_models(log_callback=None):
    global spacy, nlp_spacy, pipeline_pii, fitz
    global MODELS_LOADED

    def log(msg):
        if log_callback:
            log_callback(msg)
        print(msg)

    try:
        ensure_models_exist(log)
    except Exception:
        log("❌ Krytyczny błąd: Nie można pobrać modeli.")
        return

    os.environ['TRANSFORMERS_OFFLINE'] = '1'
    os.environ['HF_DATASETS_OFFLINE'] = '1'

    try:
        log("⏳ Ładowanie pymupdf...")
        import fitz as _fitz
        fitz = _fitz
        log("✅ pymupdf OK")
    except ImportError:
        log("❌ Błąd: pymupdf nie jest zainstalowany")

    try:
        log("⏳ Ładowanie spaCy (Offline)...")
        import spacy as _spacy
        spacy = _spacy
        path_spacy = str(BASE_MODEL_PATH / "pl_core_news_lg")
        nlp_spacy = spacy.load(path_spacy)
        log("✅ spaCy OK")
    except Exception as e:
        log(f"⚠️ spaCy niedostępny: {e}")

    try:
        from transformers import pipeline as hf_pipeline
        log("⏳ Ładowanie BardsAI PII (Offline)...")
        path_pii = str(BASE_MODEL_PATH / "eu-pii-anonimization")
        pipeline_pii = hf_pipeline(  # type: ignore[arg-type, call-overload]
            task="ner",
            model=path_pii,
            tokenizer=path_pii,
            aggregation_strategy="simple",
        )
        log("✅ BardsAI PII OK")
    except Exception as e:
        log(f"⚠️ Błąd ładowania Transformers: {e}")

    try:
        log("⏳ Ładowanie bazy imion (Gazetteer)...")
        load_names_db(log)
    except Exception as e:
        log(f"⚠️ Błąd bazy imion: {e}")

    MODELS_LOADED = True
    log("🚀 Gotowe! Wszystkie dostępne modele zostały załadowane.")


# ─────────────────────────────────────────────
# Мапування RegEx і NER
# ─────────────────────────────────────────────
REGEX_PATTERNS = {
    "pesel":       (r"\b\d{11}\b", "PESEL"),
    "nip":         (r"\b(\d{3}[-–]?\d{3}[-–]?\d{2}[-–]?\d{2}|\d{10})\b", "NIP"),
    "regon":       (r"\b(\d{9}|\d{14})\b", "REGON"),
    "iban":        (r"\b(PL\d{2}[\s]?\d{4}[\s]?\d{4}[\s]?\d{4}[\s]?\d{4}[\s]?\d{4}[\s]?\d{4})\b", "IBAN"),
    "card":        (r"\b(?:\d[ -]*?){13,16}\b", "KARTA"),
    "cvv":         (r"\b\d{3}\b(?=\s|$)", "CVV"),
    "phone":       (r"(?<!\d)(\+48[\s\-]?)?(\d{3}[\s\-]?\d{3}[\s\-]?\d{3})(?!\d)", "TELEFON"),
    "email":       (r"\b([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})\b", "EMAIL"),
    "ip":          (r"\b(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})\b", "IP"),
    "postal_code": (r"\b(\d{2}-\d{3})\b", "KOD_POCZTOWY"),
    "passport":    (r"\b([A-Z]{2}\d{7})\b", "PASZPORT"),
    "pwz":         (r"\b(PWZ[\s:]?\d{7})\b", "PWZ"),
    "street":      (r"\b(ul\.|ulica)\s+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+\b", "ADRES"),
    "chat_log":    (r"\b[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+(?:[- ][A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+){0,2}[\s\xa0]+(?:\d{1,2}:)?\d{1,3}:\d{2}\b", "LOG_CZATU"),
}

NER_LABEL_MAP = {
    "persName": "OSOBA", "PERSON": "OSOBA", "PER": "OSOBA", "nam_liv_person": "OSOBA",
    "placeName": "ADRES", "LOC": "ADRES", "GPE": "ADRES", "LOCATION": "ADRES", "nam_loc": "ADRES", "FAC": "ADRES", "ADDRESS": "ADRES",
    "orgName": "ORGANIZACJA", "ORG": "ORGANIZACJA", "nam_org": "ORGANIZACJA", "ORGANIZATION": "ORGANIZACJA",
}

CATEGORY_TO_NER = {
    "names":     ["OSOBA"],
    "addresses": ["ADRES"],
    "companies": ["ORGANIZACJA"],
}

# FIX: зворотній словник — O(1) замість O(n) next() у кожному виклику NER
MAPPED_TO_CAT: dict[str, str] = {
    label: cat
    for cat, labels in CATEGORY_TO_NER.items()
    for label in labels
}

CATEGORY_TO_REGEX = {
    "pesel": "pesel", "nip": "nip", "regon": "regon", "iban": "iban",
    "phone": "phone", "email": "email", "ip": "ip", "postal_code": "postal_code",
    "passport": "passport", "pwz": "pwz", "card": "card", "cvv": "cvv",
    "addresses": "street", "names": "chat_log",
    # NOTE: "logos" навмисно відсутній — обробляється окремо через OpenCV template matching
}


# ─────────────────────────────────────────────
# Логіка анонімізації
# ─────────────────────────────────────────────
class TokenRegistry:
    def __init__(self):
        self.registry: dict[str, dict] = {}
        self._counter: dict[str, int] = {}

    def get_or_create(self, original: str, label: str) -> str:
        key = original.strip()
        if key in self.registry:
            return self.registry[key]["token"]
        prefix = label[:3].upper()
        if label == "LOG_CZATU":
            prefix = "LOG"
        idx = self._counter.get(label, 0) + 1
        self._counter[label] = idx
        token = f"[{prefix}_{idx:03d}]"
        self.registry[key] = {"token": token, "label": label, "original": key, "count": 1}
        return token

    # ВИДАЛЕНО: get_token() — зайва обгортка навколо get_or_create

    def to_dict(self) -> dict:
        return self.registry


def run_names_db(text: str, enabled_cats: set, registry: TokenRegistry) -> list:
    findings = []
    if "names" not in enabled_cats or not KNOWN_NAMES_SET:
        return findings
    pattern = re.compile(r'\b[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+\b')
    for match in pattern.finditer(text):
        word = match.group(0)
        if word.lower() in STOP_WORDS:
            continue
        if word.lower() in KNOWN_NAMES_SET:
            # FIX: викликаємо get_or_create напряму замість видаленого get_token()
            findings.append((match.start(), match.end(), registry.get_or_create(word, "OSOBA"), "OSOBA", "regex"))
    return findings


def run_regex(text: str, enabled_cats: set, registry: TokenRegistry) -> list:
    findings = []
    for cat, pattern_key in CATEGORY_TO_REGEX.items():
        if cat not in enabled_cats:
            continue
        pattern, label = REGEX_PATTERNS[pattern_key]

        for m in re.finditer(pattern, text, re.IGNORECASE):
            if label == "CVV":
                line_start = text.rfind('\n', 0, m.start())
                line_start = 0 if line_start == -1 else line_start + 1
                line_end = text.find('\n', m.end())
                line_end = len(text) if line_end == -1 else line_end
                line_text = text[line_start:line_end]
                if not re.search(r'\b(?:cvv|cvc)\b', line_text, re.IGNORECASE):
                    continue

            token = registry.get_or_create(m.group(), label)
            findings.append((m.start(), m.end(), token, label, "regex"))

    return findings


def run_spacy(text: str, enabled_cats: set, registry: TokenRegistry) -> list:
    if nlp_spacy is None:
        return []
    findings = []
    doc = nlp_spacy(text)
    for ent in doc.ents:
        if ent.text.strip().lower() in STOP_WORDS or len(ent.text.strip()) < 2:
            continue
        mapped = NER_LABEL_MAP.get(ent.label_)
        if mapped is None:
            continue
        # FIX: використовуємо зворотній словник MAPPED_TO_CAT
        cat = MAPPED_TO_CAT.get(mapped)
        if cat and cat in enabled_cats:
            token = registry.get_or_create(ent.text, mapped)
            findings.append((ent.start_char, ent.end_char, token, mapped, "spacy"))
    return findings


def run_transformer(pipe, name: str, text: str, enabled_cats: set, registry: TokenRegistry) -> list:
    if pipe is None:
        return []
    findings = []
    try:
        results = pipe(text)
        for ent in results:
            start = ent.get("start", 0)
            end = ent.get("end", 0)
            word = text[start:end]
            if word.strip().lower() in STOP_WORDS or len(word.strip()) < 2:
                continue
            raw_label = ent.get("entity_group", ent.get("entity", ""))
            mapped = NER_LABEL_MAP.get(raw_label, NER_LABEL_MAP.get(raw_label.upper()))
            if mapped is None:
                continue
            # FIX: використовуємо зворотній словник MAPPED_TO_CAT
            cat = MAPPED_TO_CAT.get(mapped)
            if cat and cat in enabled_cats:
                token = registry.get_or_create(word, mapped)
                findings.append((start, end, token, mapped, name))
    except Exception as e:
        print(f"⚠️ Błąd {name}: {e}")
    return findings


def vote_and_anonymize(text: str, all_findings: list) -> tuple[str, list]:
    ml_count = sum(1 for s in [nlp_spacy, pipeline_pii] if s is not None)
    threshold = max(1, ml_count * 0.5)

    span_votes: dict[tuple, dict] = {}
    for start, end, token, label, source in all_findings:
        key = (start, end)
        if key not in span_votes:
            span_votes[key] = {"token": token, "label": label, "force": False, "votes": set()}
        if source == "regex":
            span_votes[key]["force"] = True
        else:
            span_votes[key]["votes"].add(source)

    to_anonymize = []
    for (start, end), info in span_votes.items():
        if info["force"] or len(info["votes"]) >= threshold:
            to_anonymize.append((start, end, info["token"], info["label"]))

    to_anonymize.sort(key=lambda x: x[0])
    merged = []
    for span in to_anonymize:
        if not merged or span[0] >= merged[-1][1]:
            merged.append(list(span))
        elif span[1] > merged[-1][1]:
            merged[-1][1] = span[1]

    result = []
    prev = 0
    for start, end, token, label in merged:
        result.append(text[prev:start])
        result.append(token)
        prev = end
    result.append(text[prev:])

    return "".join(result), merged


# ─────────────────────────────────────────────
# Logo Selector (OpenCV + Tkinter)
# ─────────────────────────────────────────────
class LogoSelector(tk.Toplevel):
    def __init__(self, parent, pdf_path: str):
        super().__init__(parent)
        self.title("Wybierz logo (Zaznacz i zamknij okno)")
        self.geometry("800x900")

        self.pdf_path = pdf_path
        self.template_image = None
        self.zoom = 2.0

        self.rect_id: int | None = None          # FIX: явна анотація типу
        self.start_x: float = 0.0               # FIX: ініціалізація як float
        self.start_y: float = 0.0

        self._load_first_page()
        self._build_canvas()

    def _load_first_page(self):
        if fitz is None:
            raise RuntimeError("pymupdf nie jest załadowany")  # FIX: guard для fitz
        doc = fitz.open(self.pdf_path)
        page = doc[0]
        mat = fitz.Matrix(self.zoom, self.zoom)  # тепер fitz точно не None
        self.pix = page.get_pixmap(matrix=mat)
        # FIX: tuple замість list для Image.frombytes
        img = Image.frombytes("RGB", (self.pix.width, self.pix.height), self.pix.samples)
        self.tk_image = ImageTk.PhotoImage(img)
        doc.close()

    def _build_canvas(self):
        self.canvas = tk.Canvas(self, cursor="cross")
        self.canvas.pack(fill="both", expand=True)
        self.canvas.create_image(0, 0, anchor="nw", image=self.tk_image)

        self.canvas.bind("<ButtonPress-1>", self._on_press)
        self.canvas.bind("<B1-Motion>", self._on_drag)
        self.canvas.bind("<ButtonRelease-1>", self._on_release)

    def _on_press(self, event):
        self.start_x = self.canvas.canvasx(event.x)
        self.start_y = self.canvas.canvasy(event.y)
        if self.rect_id is not None:
            self.canvas.delete(self.rect_id)
        self.rect_id = self.canvas.create_rectangle(
            self.start_x, self.start_y, self.start_x, self.start_y,
            outline="red", width=3,
        )

    def _on_drag(self, event):
        if self.rect_id is None:  # FIX: guard проти None
            return
        cur_x = self.canvas.canvasx(event.x)
        cur_y = self.canvas.canvasy(event.y)
        self.canvas.coords(self.rect_id, self.start_x, self.start_y, cur_x, cur_y)

    def _on_release(self, event):
        end_x = self.canvas.canvasx(event.x)
        end_y = self.canvas.canvasy(event.y)

        # FIX: start_x/start_y тепер завжди float — min/max не падають
        x1, y1 = min(self.start_x, end_x), min(self.start_y, end_y)
        x2, y2 = max(self.start_x, end_x), max(self.start_y, end_y)

        if abs(x2 - x1) > 10 and abs(y2 - y1) > 10:
            img_np = np.frombuffer(self.pix.samples, dtype=np.uint8).reshape(
                self.pix.height, self.pix.width, self.pix.n
            )
            if self.pix.n == 4:
                img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGBA2BGR)
            else:
                img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)
            self.template_image = img_cv[int(y1):int(y2), int(x1):int(x2)]
            messagebox.showinfo("Zapisano", "Logo zostało skopiowane jako wzór. Możesz zamknąć to okno.")


# ─────────────────────────────────────────────
# Обробка файлів
# ─────────────────────────────────────────────
def anonymize_pdf(
    input_path: str,
    output_path: str,
    enabled_cats: set,
    registry: TokenRegistry,
    log_cb=None,
    logo_template=None,
) -> dict:
    if fitz is None:
        if log_cb:
            log_cb("❌ Błąd: pymupdf nie jest zainstalowany")
        return {}

    doc = fitz.open(input_path)

    # FIX: ітерація через range(len(doc)) — Document не реалізує __iter__ для enumerate
    for page_num in range(len(doc)):
        page = doc[page_num]
        if log_cb:
            log_cb(f"📄 Przetwarzanie strony {page_num + 1}/{len(doc)}...")

        text: str = str(page.get_text("text"))  # FIX: явний cast до str
        search_page = page
        ocr_doc = None
        draw_queue: list = []  # оголошується один раз на початку ітерації

        # 1. ТЕМПЛЕЙТ МЕТЧІНГ (ПОШУК ЛОГОТИПІВ)
        if logo_template is not None:
            try:
                zoom = 2.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)

                img_np = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
                if pix.n == 4:
                    img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGBA2BGR)
                else:
                    img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)

                result = cv2.matchTemplate(img_cv, logo_template, cv2.TM_CCOEFF_NORMED)
                threshold = 0.8
                locations = np.where(result >= threshold)

                h, w = logo_template.shape[:2]
                detected_count = 0

                for pt in zip(*locations[::-1]):
                    x1, y1 = pt[0], pt[1]
                    x2, y2 = pt[0] + w, pt[1] + h
                    pdf_rect = fitz.Rect(x1 / zoom, y1 / zoom, x2 / zoom, y2 / zoom)
                    page.add_redact_annot(pdf_rect, fill=(0, 0, 0))

                    token = "[LOGO]"
                    font_size = max(4, min(8, pdf_rect.height * 0.7))
                    text_width = fitz.get_text_length(token, fontsize=font_size)
                    t_x = pdf_rect.x0 + (pdf_rect.width - text_width) / 2
                    t_y = pdf_rect.y1 - (pdf_rect.height - font_size) / 2
                    draw_queue.append((t_x, t_y, token, font_size))
                    detected_count += 1

                if detected_count > 0 and log_cb:
                    log_cb(f"🏢 Znaleziono logo w {detected_count} miejscach.")
            except Exception as e:
                if log_cb:
                    log_cb(f"⚠️ Błąd wyszukiwania logo: {e}")

        # 2. ДЕТЕКЦІЯ ТЕКСТУ (OCR)
        if len(text.strip()) < 50:
            if log_cb:
                log_cb(f"🔍 Wykryto skan (strona {page_num + 1}). Uruchamiam zaawansowany OCR...")
            # FIX: old_cwd ініціалізується до try, щоб завжди був bound у except
            old_cwd = os.getcwd()
            try:
                tess_dir = str((BASE_MODEL_PATH / "tesseract").absolute())
                tessdata_dir = str((BASE_MODEL_PATH / "tesseract" / "tessdata").absolute())

                os.chdir(tess_dir)
                pix = page.get_pixmap(dpi=300)
                ocr_pdf_bytes: bytes = bytes(pix.pdfocr_tobytes(language="pol", tessdata=tessdata_dir))  # type: ignore[arg-type]
                os.chdir(old_cwd)

                ocr_doc = fitz.open("pdf", ocr_pdf_bytes)
                search_page = ocr_doc[0]
                text = str(search_page.get_text("text"))  # FIX: явний cast

                if log_cb:
                    log_cb(f"📝 OCR wyciągnął: {len(text)} znaków.")
            except Exception as e:
                os.chdir(old_cwd)  # тепер завжди bound
                if log_cb:
                    log_cb(f"⚠️ Błąd OCR: {e}")

        if not text.strip():
            if ocr_doc:
                ocr_doc.close()
            if draw_queue:
                page.apply_redactions()
                for x, y, tkn, fs in draw_queue:
                    page.insert_text((x, y), tkn, fontsize=fs, color=(1, 1, 1))
            continue

        all_findings: list = []
        all_findings.extend(run_names_db(text, enabled_cats, registry))
        all_findings.extend(run_regex(text, enabled_cats, registry))
        all_findings.extend(run_spacy(text, enabled_cats, registry))
        all_findings.extend(run_transformer(pipeline_pii, "bardsai", text, enabled_cats, registry))

        _, merged = vote_and_anonymize(text, all_findings)

        for start, end, token, label in merged:
            word = text[start:end].strip()
            if not word or len(word) < 2:
                continue
            instances = search_page.search_for(word)
            for rect in instances:
                page.add_redact_annot(rect, fill=(0, 0, 0))
                font_size = max(4, min(8, rect.height * 0.7))
                text_width = fitz.get_text_length(token, fontsize=font_size)
                t_x = rect.x0 + (rect.width - text_width) / 2
                t_y = rect.y1 - (rect.height - font_size) / 2
                draw_queue.append((t_x, t_y, token, font_size))

        page.apply_redactions()

        for x, y, tkn, fs in draw_queue:
            page.insert_text((x, y), tkn, fontsize=fs, color=(1, 1, 1))

        # FIX: ocr_doc закривається один раз — через finally-подібний патерн в кінці ітерації
        if ocr_doc:
            ocr_doc.close()
            ocr_doc = None

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    if log_cb:
        log_cb(f"✅ PDF zapisano: {output_path}")
    return registry.to_dict()


def anonymize_docx(
    input_path: str,
    output_path: str,
    enabled_cats: set,
    registry: TokenRegistry,
    log_cb=None,
) -> dict:
    try:
        import docx
    except ImportError:
        if log_cb:
            log_cb("❌ Błąd: biblioteka 'python-docx' nie jest zainstalowana.")
        return {}

    if log_cb:
        log_cb("📄 Przetwarzanie dokumentu Word...")
    doc = docx.Document(input_path)

    def process_text_block(block):
        if not block.text.strip():
            return
        text: str = block.text
        all_findings: list = []
        all_findings.extend(run_names_db(text, enabled_cats, registry))
        all_findings.extend(run_regex(text, enabled_cats, registry))
        all_findings.extend(run_spacy(text, enabled_cats, registry))
        all_findings.extend(run_transformer(pipeline_pii, "bardsai", text, enabled_cats, registry))
        anonymized_text, merged = vote_and_anonymize(text, all_findings)
        if merged:
            block.text = anonymized_text

    for para in doc.paragraphs:
        process_text_block(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_text_block(para)

    doc.save(output_path)
    if log_cb:
        log_cb(f"✅ Plik DOCX zapisano: {output_path}")
    return registry.to_dict()


def anonymize_text_file(
    input_path: str,
    output_path: str,
    enabled_cats: set,
    registry: TokenRegistry,
    log_cb=None,
) -> dict:
    with open(input_path, "r", encoding="utf-8", errors="replace") as f:
        text: str = f.read()

    all_findings: list = []
    all_findings.extend(run_names_db(text, enabled_cats, registry))
    all_findings.extend(run_regex(text, enabled_cats, registry))
    all_findings.extend(run_spacy(text, enabled_cats, registry))
    all_findings.extend(run_transformer(pipeline_pii, "bardsai", text, enabled_cats, registry))

    anonymized_text, _ = vote_and_anonymize(text, all_findings)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(anonymized_text)

    if log_cb:
        log_cb(f"✅ Plik zapisano: {output_path}")
    return registry.to_dict()


class HotFolderHandler(FileSystemEventHandler):
    def __init__(self, file_queue: queue.Queue, log_cb):
        self.queue = file_queue
        self.log_cb = log_cb

    def on_created(self, event):
        if event.is_directory:
            return
        raw_path = event.src_path
        # FIX: watchdog може повернути bytes на деяких ОС
        path: str = raw_path.decode("utf-8") if isinstance(raw_path, bytes) else raw_path
        if "~$" not in path and any(path.lower().endswith(ext) for ext in ['.pdf', '.docx', '.txt', '.text']):
            self.log_cb(f"📥 Wykryto nowy plik w folderze: {Path(path).name}")
            self.queue.put(path)


# ─────────────────────────────────────────────
# GUI (Light Blue Theme)
# ─────────────────────────────────────────────
DARK_BG   = "#e0f2fe"
PANEL_BG  = "#bae6fd"
CARD_BG   = "#ffffff"
ACCENT    = "#0284c7"
ACCENT2   = "#0369a1"
TEXT_PRIMARY   = "#0f172a"
TEXT_SECONDARY = "#334155"
GREEN  = "#16a34a"
RED    = "#dc2626"
YELLOW = "#d6da16"
BORDER = "#7dd3fc"


class AnonymizerApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("🔒 Polish Document Anonymizer (Wersja Offline)")
        self.geometry("980x740")
        self.configure(bg=DARK_BG)
        self.resizable(True, True)

        self.selected_file = tk.StringVar()
        self.status_var = tk.StringVar(value="Gotowy")
        self.registry = TokenRegistry()

        self.cat_vars: dict[str, tk.BooleanVar] = {
            "names":       tk.BooleanVar(value=True),
            "addresses":   tk.BooleanVar(value=True),
            "companies":   tk.BooleanVar(value=True),
            "pesel":       tk.BooleanVar(value=True),
            "nip":         tk.BooleanVar(value=True),
            "regon":       tk.BooleanVar(value=True),
            "iban":        tk.BooleanVar(value=True),
            "card":        tk.BooleanVar(value=True),
            "cvv":         tk.BooleanVar(value=True),
            "phone":       tk.BooleanVar(value=True),
            "email":       tk.BooleanVar(value=True),
            "ip":          tk.BooleanVar(value=True),
            "postal_code": tk.BooleanVar(value=True),
            "passport":    tk.BooleanVar(value=True),
            "pwz":         tk.BooleanVar(value=True),
            "logos":       tk.BooleanVar(value=True),
        }

        self.watch_src = tk.StringVar()
        self.watch_dst = tk.StringVar()
        self.is_watching = False
        self.observer: "Optional[Observer]" = None  # FIX: явна анотація типу
        self.auto_queue: queue.Queue = queue.Queue()

        self.last_activity = time.time()
        self.is_processing = False
        self.tray_icon = None

        self._build_ui()
        self._start_model_loading()

        self.protocol("WM_DELETE_WINDOW", self._hide_to_tray)
        threading.Thread(target=self._idle_monitor, daemon=True).start()
        threading.Thread(target=self._auto_process_worker, daemon=True).start()

    def _toggle_autostart(self):
        try:
            import winreg
            key_path = r"Software\Microsoft\Windows\CurrentVersion\Run"
            app_name = "AnonSupreme"
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path, 0, winreg.KEY_ALL_ACCESS)
            try:
                winreg.QueryValueEx(key, app_name)
                winreg.DeleteValue(key, app_name)
                messagebox.showinfo("Autostart", "❌ Program został usunięty z autostartu Windows.")
            except FileNotFoundError:
                winreg.SetValueEx(key, app_name, 0, winreg.REG_SZ, f'"{EXE_PATH}"')
                messagebox.showinfo("Autostart", "✅ Program został dodany do autostartu Windows.")
            finally:
                winreg.CloseKey(key)
        except Exception as e:
            self._log(f"❌ Błąd autostartu: {e}")
            messagebox.showerror("Błąd", f"Nie udało się zmienić ustawień:\n{e}")

    def _build_ui(self):
        header = tk.Frame(self, bg=DARK_BG)
        header.pack(fill="x", padx=20, pady=(18, 0))

        tk.Label(
            header, text="🔒 Polish Document Anonymizer",
            font=("Courier New", 18, "bold"), bg=DARK_BG, fg=ACCENT,
        ).pack(side="left")
        self.model_status_label = tk.Label(
            header, text="● Ładowanie modeli...",
            font=("Courier New", 9, "bold"), bg=DARK_BG, fg=YELLOW,
        )
        self.model_status_label.pack(side="right", padx=10)

        content = tk.Frame(self, bg=DARK_BG)
        content.pack(fill="both", expand=True, padx=20, pady=12)

        left = tk.Frame(content, bg=DARK_BG)
        left.pack(side="left", fill="both", expand=True)

        right = tk.Frame(content, bg=DARK_BG, width=320)
        right.pack(side="right", fill="y", padx=(12, 0))
        right.pack_propagate(False)

        file_card = self._card(left, "📂 Wybór dokumentu")
        file_row = tk.Frame(file_card, bg=CARD_BG)
        file_row.pack(fill="x")
        self.file_entry = tk.Entry(
            file_row, textvariable=self.selected_file,
            font=("Courier New", 9), bg="#f0f9ff", fg=TEXT_PRIMARY,
            insertbackground=ACCENT, relief="flat", bd=6,
        )
        self.file_entry.pack(side="left", fill="x", expand=True)
        self._btn(file_row, "Przeglądaj", self._browse_file, ACCENT).pack(side="right", padx=(8, 0))

        cat_card = self._card(left, "🎯 Kategorie do anonimizacji")
        cat_labels = {
            "names":       "👤 Imiona i nazwiska",
            "addresses":   "📍 Adresy i miasta",
            "companies":   "🏢 Firmy i org.",
            "pesel":       "🆔 PESEL",
            "nip":         "🏦 NIP",
            "regon":       "📋 REGON",
            "iban":        "💳 IBAN / Konto",
            "card":        "💳 Numer karty",
            "cvv":         "🔐 CVV",
            "phone":       "📱 Telefony",
            "email":       "✉️ Adresy E-mail",
            "ip":          "🌐 Adresy IP",
            "postal_code": "📮 Kody pocztowe",
            "passport":    "🛂 Paszporty",
            "pwz":         "🏥 PWZ (lekarze)",
            "logos":       "🏢 Logotypy firm",
        }

        grid_frame = tk.Frame(cat_card, bg=CARD_BG)
        grid_frame.pack(fill="x")

        cols = 3
        for i, (key, label) in enumerate(cat_labels.items()):
            row_i, col_i = i // cols, i % cols
            cb = tk.Checkbutton(
                grid_frame, text=label, variable=self.cat_vars[key],
                bg=CARD_BG, fg=TEXT_PRIMARY, selectcolor="#f0f9ff",
                activebackground=CARD_BG, activeforeground=ACCENT,
                font=("Courier New", 8), cursor="hand2",
            )
            cb.grid(row=row_i, column=col_i, sticky="w", padx=6, pady=2)

        btn_row = tk.Frame(cat_card, bg=CARD_BG)
        btn_row.pack(fill="x", pady=(6, 0))
        self._btn(btn_row, "Zaznacz wszystko", self._select_all, ACCENT, small=True).pack(side="left")
        self._btn(btn_row, "Odznacz wszystko", self._deselect_all, RED, small=True).pack(side="left", padx=6)

        auto_card = self._card(left, "🔄 Tryb automatyczny (Hot Folder)")
        src_row = tk.Frame(auto_card, bg=CARD_BG)
        src_row.pack(fill="x", pady=(0, 4))
        tk.Button(
            src_row, text="📁 Źródło", command=self._browse_src_dir,
            font=("Courier New", 8, "bold"), bg=ACCENT, fg="white", relief="flat", padx=6,
        ).pack(side="left")
        tk.Entry(
            src_row, textvariable=self.watch_src,
            font=("Courier New", 8), bg="#f8fafc", fg=TEXT_SECONDARY,
            relief="flat", state="readonly",
        ).pack(side="left", fill="x", expand=True, padx=(6, 0))

        dst_row = tk.Frame(auto_card, bg=CARD_BG)
        dst_row.pack(fill="x", pady=(4, 8))
        tk.Button(
            dst_row, text="📁 Zapis ", command=self._browse_dst_dir,
            font=("Courier New", 8, "bold"), bg=ACCENT, fg="white", relief="flat", padx=6,
        ).pack(side="left")
        tk.Entry(
            dst_row, textvariable=self.watch_dst,
            font=("Courier New", 8), bg="#f8fafc", fg=TEXT_SECONDARY,
            relief="flat", state="readonly",
        ).pack(side="left", fill="x", expand=True, padx=(6, 0))

        self.watch_btn = self._btn(auto_card, "▶ Uruchom nasłuchiwanie", self._toggle_watch, GREEN)
        self.watch_btn.pack(fill="x")

        self.run_btn = self._btn(left, "🚀  ANONIMIZUJ", self._run_anonymization, ACCENT2, large=True)
        self.run_btn.pack(fill="x", pady=(10, 0))

        log_card = self._card(left, "📋 Log operacji")
        self.log_text = scrolledtext.ScrolledText(
            log_card, height=10, font=("Courier New", 8),
            bg="#f8fafc", fg=GREEN, insertbackground=GREEN,
            relief="flat", bd=1, state="disabled",
        )
        self.log_text.pack(fill="both", expand=True)

        reg_label = tk.Label(
            right, text="🗃️  Rejestr tokenów",
            font=("Courier New", 10, "bold"), bg=DARK_BG, fg=TEXT_PRIMARY,
        )
        reg_label.pack(anchor="w", pady=(0, 6))
        self.registry_text = scrolledtext.ScrolledText(
            right, font=("Courier New", 8), bg="#ffffff", fg=TEXT_PRIMARY,
            insertbackground=ACCENT, relief="flat", bd=1, state="disabled",
            highlightbackground=BORDER, highlightthickness=1,
        )
        self.registry_text.pack(fill="both", expand=True)

        export_row = tk.Frame(right, bg=DARK_BG)
        export_row.pack(fill="x", pady=(6, 0))
        self._btn(export_row, "💾 Eksport JSON", self._export_registry, ACCENT, small=True).pack(fill="x", pady=(0, 6))
        self._btn(export_row, "⚙️ Dodaj do menu (Prawy klik)", self._add_context_menu, ACCENT2, small=True).pack(fill="x")
        self._btn(export_row, "🚀 Autostart z Windows", self._toggle_autostart, ACCENT, small=True).pack(fill="x")

        status_bar = tk.Frame(self, bg=PANEL_BG, height=24)
        status_bar.pack(fill="x", side="bottom")
        tk.Label(
            status_bar, textvariable=self.status_var,
            font=("Courier New", 8), bg=PANEL_BG, fg=TEXT_SECONDARY,
        ).pack(side="left", padx=10)

    def _card(self, parent, title: str) -> tk.Frame:
        wrapper = tk.Frame(parent, bg=DARK_BG)
        wrapper.pack(fill="x", pady=(0, 10))
        tk.Label(wrapper, text=title, font=("Courier New", 9, "bold"), bg=DARK_BG, fg=TEXT_SECONDARY).pack(anchor="w", pady=(0, 4))
        card = tk.Frame(wrapper, bg=CARD_BG, bd=0, relief="flat", padx=12, pady=10)
        card.pack(fill="x")
        card.configure(highlightbackground=BORDER, highlightthickness=1)
        return card

    def _btn(self, parent, text: str, command, color: str, small=False, large=False) -> tk.Button:
        size = 8 if small else (12 if large else 9)
        pady = 2 if small else (8 if large else 4)
        return tk.Button(
            parent, text=text, command=command,
            font=("Courier New", size, "bold"), bg=color, fg="white",
            activebackground=PANEL_BG, activeforeground=color,
            relief="flat", bd=0, cursor="hand2", pady=pady, padx=10,
        )

    def _browse_file(self):
        path = filedialog.askopenfilename(
            title="Wybierz dokument",
            filetypes=[
                ("Obsługiwane pliki", "*.pdf *.docx *.txt *.text"),
                ("PDF", "*.pdf"),
                ("Dokument Word", "*.docx"),
                ("Pliki tekstowe", "*.txt *.text"),
            ],
        )
        if path:
            self.selected_file.set(path)

    def _select_all(self):
        for v in self.cat_vars.values():
            v.set(True)

    def _deselect_all(self):
        for v in self.cat_vars.values():
            v.set(False)

    def _log(self, msg: str):
        def _do():
            self.log_text.configure(state="normal")
            ts = datetime.now().strftime("%H:%M:%S")
            self.log_text.insert("end", f"[{ts}] {msg}\n")
            self.log_text.see("end")
            self.log_text.configure(state="disabled")
        self.after(0, _do)

    def _update_registry_panel(self):
        def _do():
            reg = self.registry.to_dict()
            self.registry_text.configure(state="normal")
            self.registry_text.delete("1.0", "end")
            if not reg:
                self.registry_text.insert("end", "Pusto\n")
            else:
                for original, info in sorted(reg.items(), key=lambda x: x[1]["token"]):
                    line = f"{info['token']:12s} │ {info['label']:12s} │ {original[:30]}\n"
                    self.registry_text.insert("end", line)
            self.registry_text.configure(state="disabled")
        self.after(0, _do)

    def _export_registry(self):
        if not self.registry.to_dict():
            messagebox.showinfo("Uwaga", "Rejestr jest pusty")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON", "*.json")],
            initialfile="token_registry.json",
        )
        if path:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.registry.to_dict(), f, ensure_ascii=False, indent=2)
            messagebox.showinfo("✅ Zapisano", f"Rejestr zapisany:\n{path}")

    def _start_model_loading(self):
        def _done_cb(msg):
            if "Gotowe" in msg:
                self.after(0, lambda: self.model_status_label.configure(text="● Modele gotowe", fg=GREEN))
            self._log(msg)
        threading.Thread(target=lazy_load_models, args=(_done_cb,), daemon=True).start()

    def _add_context_menu(self):
        try:
            import winreg
            # FIX: завжди пишемо в HKEY_CURRENT_USER (не потребує прав адміністратора)
            key_path = r"Software\Classes\*\shell\Anonimizuj"
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path)
            winreg.SetValue(key, "", winreg.REG_SZ, "🔒 Anonimizuj dokument")
            winreg.SetValueEx(key, "Icon", 0, winreg.REG_SZ, f'"{EXE_PATH}"')
            winreg.CloseKey(key)

            cmd_key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path + r"\command")
            winreg.SetValue(cmd_key, "", winreg.REG_SZ, f'"{EXE_PATH}" "%1"')
            winreg.CloseKey(cmd_key)
            self._log("✅ Zaktualizowano menu kontekstowe Windows.")
            messagebox.showinfo("✅ Sukces", "Opcja dodana do menu pod prawym przyciskiem myszy!")
        except Exception as e:
            err_msg = str(e)
            self._log(f"❌ Błąd rejestru: {err_msg}")
            messagebox.showerror("❌ Błąd", f"Nie udało się zaktualizować rejestru:\n{err_msg}")

    def _run_anonymization(self):
        path = self.selected_file.get().strip()
        if not path or not os.path.exists(path):
            messagebox.showwarning("⚠️ Uwaga", "Najpierw wybierz istniejący dokument!")
            return

        enabled_cats = {k for k, v in self.cat_vars.items() if v.get()}
        if not enabled_cats:
            messagebox.showwarning("⚠️ Uwaga", "Wybierz przynajmniej jedną kategorię!")
            return

        p = Path(path)
        out_path = filedialog.asksaveasfilename(
            title="Zapisz zanonimizowany plik",
            initialfile=f"{p.stem}_anonymized{p.suffix}",
            defaultextension=p.suffix,
            filetypes=[("Taki sam format", f"*{p.suffix}")],
        )
        if not out_path:
            return

        self.run_btn.configure(state="disabled", text="⏳ Przetwarzanie...")
        self.registry = TokenRegistry()

        logo_template_cv = None
        if path.lower().endswith(".pdf") and "logos" in enabled_cats:
            selector = LogoSelector(self, path)
            self.wait_window(selector)
            if selector.template_image is not None:
                logo_template_cv = selector.template_image
                self._log("✅ Pobrano wzór logotypu.")
            else:
                self._log("⚠️ Nie zaznaczono logotypu, pomijam ten krok.")

        def worker(cv_template):
            self.is_processing = True
            self.last_activity = time.time()

            global MODELS_LOADED
            if not MODELS_LOADED:
                self._log("⏳ Wybudzanie modeli AI...")
                self.after(0, lambda: self.model_status_label.configure(text="● Ładowanie modeli...", fg=YELLOW))
                lazy_load_models(self._log)
                self.after(0, lambda: self.model_status_label.configure(text="● Modele gotowe", fg=GREEN))

            self._log(f"🔒 Rozpoczęcie anonimizacji: {p.name}")
            try:
                if path.lower().endswith(".pdf"):
                    anonymize_pdf(path, out_path, enabled_cats, self.registry, self._log, cv_template)
                elif path.lower().endswith(".docx"):
                    anonymize_docx(path, out_path, enabled_cats, self.registry, self._log)
                else:
                    anonymize_text_file(path, out_path, enabled_cats, self.registry, self._log)

                out_p = Path(out_path)
                reg_path = out_p.parent / f"{out_p.stem}_tokens.json"
                with open(reg_path, "w", encoding="utf-8") as f:
                    json.dump(self.registry.to_dict(), f, ensure_ascii=False, indent=2)

                self._update_registry_panel()
                reg_count = len(self.registry.to_dict())
                self._log(f"🎉 Gotowe! Zapisano dokument i {reg_count} tokenów.")
                self.after(0, lambda c=reg_count: messagebox.showinfo(
                    "✅ Gotowe", "Anonimizacja zakończona!\nDokument i rejestr JSON zapisane."
                ))
            except Exception as e:
                err_msg = str(e)
                self._log(f"❌ Błąd: {err_msg}")
                self.after(0, lambda msg=err_msg: messagebox.showerror("❌ Błąd", msg))
            finally:
                self.is_processing = False
                self.last_activity = time.time()
                self.after(0, lambda: self.run_btn.configure(state="normal", text="🚀  ANONIMIZUJ"))
                self.after(0, lambda: self.status_var.set(f"Zakończono: {p.name}"))

        threading.Thread(target=worker, args=(logo_template_cv,), daemon=True).start()

    def _browse_src_dir(self):
        path = filedialog.askdirectory(title="Wybierz folder źródłowy (nasłuchiwanie)")
        if path:
            self.watch_src.set(path)

    def _browse_dst_dir(self):
        path = filedialog.askdirectory(title="Wybierz folder docelowy (zapis)")
        if path:
            self.watch_dst.set(path)

    def _toggle_watch(self):
        if self.is_watching:
            self.is_watching = False
            # FIX: guard проти None перед викликом .stop()/.join()
            if self.observer is not None:
                self.observer.stop()
                self.observer.join()
                self.observer = None
            self.watch_btn.configure(text="▶ Uruchom nasłuchiwanie", bg=GREEN)
            self._log("⏹️ Zatrzymano automatyczne nasłuchiwanie.")
            self.status_var.set("Nasłuchiwanie wyłączone")
        else:
            src, dst = self.watch_src.get(), self.watch_dst.get()
            if not src or not dst:
                messagebox.showwarning("Uwaga", "Wybierz folder źródłowy i docelowy!")
                return
            if src == dst:
                messagebox.showwarning("Uwaga", "Folder źródłowy i docelowy nie mogą być tym samym folderem!")
                return
            if not MODELS_LOADED:
                messagebox.showwarning("Uwaga", "Poczekaj na załadowanie modeli AI!")
                return

            self.is_watching = True
            self.watch_btn.configure(text="⏸️ Zatrzymaj nasłuchiwanie", bg=RED)
            self._log(f"▶️ Start nasłuchiwania w:\n{src}")
            self.status_var.set("Nasłuchiwanie aktywne...")

            event_handler = HotFolderHandler(self.auto_queue, self._log)
            self.observer = Observer()
            assert self.observer is not None  # допомагає Pylance звузити тип
            self.observer.schedule(event_handler, src, recursive=False)
            self.observer.start()

    def _auto_process_worker(self):
        while True:
            filepath: str = self.auto_queue.get()
            self.is_processing = True
            self.last_activity = time.time()
            time.sleep(2)

            global MODELS_LOADED
            if not MODELS_LOADED:
                self._log("⏳ Wybudzanie modeli AI (Auto)...")
                self.after(0, lambda: self.model_status_label.configure(text="● Ładowanie modeli...", fg=YELLOW))
                lazy_load_models(self._log)
                self.after(0, lambda: self.model_status_label.configure(text="● Modele gotowe", fg=GREEN))

            try:
                p = Path(filepath)
                dst_folder = Path(self.watch_dst.get())
                out_path = str(dst_folder / f"{p.stem}_anonymized{p.suffix}")
                reg_path = dst_folder / f"{p.stem}_tokens.json"

                enabled_cats = {k for k, v in self.cat_vars.items() if v.get()}
                self._log(f"⚙️ Auto-Przetwarzanie: {p.name}")

                if p.suffix.lower() == ".pdf":
                    anonymize_pdf(filepath, out_path, enabled_cats, self.registry, self._log)
                elif p.suffix.lower() == ".docx":
                    anonymize_docx(filepath, out_path, enabled_cats, self.registry, self._log)
                else:
                    anonymize_text_file(filepath, out_path, enabled_cats, self.registry, self._log)

                with open(reg_path, "w", encoding="utf-8") as f:
                    json.dump(self.registry.to_dict(), f, ensure_ascii=False, indent=2)

                self._update_registry_panel()
                self._log(f"✅ Auto-Zapisano: {Path(out_path).name} i rejestr JSON.")
            except Exception as e:
                self._log(f"❌ Auto-Błąd dla {Path(filepath).name}: {str(e)}")
            finally:
                self.is_processing = False
                self.last_activity = time.time()
                self.auto_queue.task_done()

    def _idle_monitor(self):
        global nlp_spacy, pipeline_pii, MODELS_LOADED
        while True:
            time.sleep(10)
            if MODELS_LOADED and not self.is_processing:
                if time.time() - self.last_activity > 300:
                    self._log("💤 5 min bezczynności. Zwalnianie pamięci RAM...")
                    nlp_spacy = None
                    pipeline_pii = None
                    gc.collect()
                    MODELS_LOADED = False
                    self.after(0, lambda: self.model_status_label.configure(
                        text="● Modele uśpione (RAM zwolniony)", fg=YELLOW
                    ))

    def _hide_to_tray(self):
        self.withdraw()
        if not self.tray_icon:
            img = Image.new('RGB', (64, 64), color=(2, 132, 199))
            d = ImageDraw.Draw(img)
            d.rectangle((16, 16, 48, 48), fill="white")
            menu = pystray.Menu(
                pystray.MenuItem("Pokaż okno", self._show_window),
                pystray.MenuItem("Zamknij całkowicie", self._quit_app),
            )
            self.tray_icon = pystray.Icon("AnonSupreme", img, "AnonSupreme - Aktywny", menu)
            threading.Thread(target=self.tray_icon.run, daemon=True).start()

    def _show_window(self, icon, item):
        if self.tray_icon is not None:
            self.tray_icon.stop()
        self.tray_icon = None
        self.after(0, self.deiconify)

    def _quit_app(self, icon, item):
        if self.tray_icon is not None:
            self.tray_icon.stop()
        os._exit(0)


# ─────────────────────────────────────────────
# Точка входу (CLI / GUI)
# ─────────────────────────────────────────────
if __name__ == "__main__":
    # FIX: блок --setup тепер пише в HKEY_CURRENT_USER (як _add_context_menu)
    # щоб не вимагати прав адміністратора
    if "--setup" in sys.argv:
        import winreg
        root = tk.Tk()
        root.withdraw()
        messagebox.showinfo(
            "Instalacja",
            "Rozpoczynam pobieranie modeli AI (ok. 1 GB).\n"
            "Może to potrwać kilka minut w zależności od połączenia sieciowego.\n\n"
            "Kliknij OK i poczekaj na komunikat końcowy.",
        )
        try:
            ensure_models_exist()
            key_path = r"Software\Classes\*\shell\Anonimizuj"  # FIX: CURRENT_USER, не CLASSES_ROOT
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path)
            winreg.SetValue(key, "", winreg.REG_SZ, "🔒 Anonimizuj dokument")
            winreg.SetValueEx(key, "Icon", 0, winreg.REG_SZ, f'"{EXE_PATH}"')
            winreg.CloseKey(key)

            cmd_key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path + r"\command")
            winreg.SetValue(cmd_key, "", winreg.REG_SZ, f'"{EXE_PATH}" "%1"')
            winreg.CloseKey(cmd_key)
            messagebox.showinfo(
                "Sukces",
                "Instalacja zakończona!\nModele zostały pobrane, a menu kontekstowe zaktualizowane.",
            )
        except Exception as e:
            messagebox.showerror("Błąd Instalacji", f"Wystąpił błąd:\n{e}")
        sys.exit(0)

    elif len(sys.argv) > 1:
        input_file = sys.argv[1]
        if not os.path.exists(input_file):
            sys.exit(1)

        p = Path(input_file)
        out_path = str(p.parent / f"{p.stem}_anonymized{p.suffix}")
        reg_path = str(p.parent / f"{p.stem}_tokens.json")

        enabled_cats = {
            "names", "addresses", "companies", "pesel", "nip", "regon",
            "iban", "card", "cvv", "phone", "email", "ip", "postal_code",
            "passport", "pwz", "logos",
        }

        registry = TokenRegistry()
        root = tk.Tk()
        root.withdraw()

        logo_template_cv = None
        if input_file.lower().endswith(".pdf"):
            if messagebox.askyesno("Wykrywanie Logo", "Czy chcesz zaznaczyć i ukryć logo w tym dokumencie?"):
                selector = LogoSelector(root, input_file)
                root.wait_window(selector)
                if selector.template_image is not None:
                    logo_template_cv = selector.template_image

        try:
            lazy_load_models()
            if input_file.lower().endswith(".pdf"):
                anonymize_pdf(input_file, out_path, enabled_cats, registry, log_cb=None, logo_template=logo_template_cv)
            elif input_file.lower().endswith(".docx"):
                anonymize_docx(input_file, out_path, enabled_cats, registry)
            else:
                anonymize_text_file(input_file, out_path, enabled_cats, registry)

            with open(reg_path, "w", encoding="utf-8") as f:
                json.dump(registry.to_dict(), f, ensure_ascii=False, indent=2)

            messagebox.showinfo("✅ Sukces", f"Anonimizacja zakończona!\n\nPlik: {Path(out_path).name}")
        except Exception as e:
            messagebox.showerror("❌ Błąd", f"Wystąpił błąd podczas anonimizacji:\n{str(e)}")
        sys.exit(0)

    else:
        app = AnonymizerApp()
        app.mainloop()